"""
ED Module Business Workflow – Word DOCX flowchart
ONLY native Word DrawingML shapes. Zero images.

Includes Marketing <-> Costing price review loop.
"""
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn as DQN
from lxml import etree
import os

_NS = {
    'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
    'mc':  'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'w':   'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp':  'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
}
def Q(ns, tag): return f'{{{_NS[ns]}}}{tag}'
def emu(n):     return int(n * 914400)

_ID = [1]
def nid(): v=_ID[0]; _ID[0]+=1; return v

# ─────────────────────────────────────────────────────────────────────────────
def _shape(x, y, w, h, prst, fill, border, bw, dashed,
           line1, line2, fs, tc, flipV, flipH, adj):
    eid = nid()
    mc  = etree.Element(Q('mc','AlternateContent'))
    ch  = etree.SubElement(mc,  Q('mc','Choice')); ch.set('Requires','wps')
    drw = etree.SubElement(ch,  Q('w','drawing'))
    anc = etree.SubElement(drw, Q('wp','anchor'))
    for k,v in [('distT','0'),('distB','0'),('distL','0'),('distR','0'),
                ('simplePos','0'),('relativeHeight',str(eid*100+3000)),
                ('behindDoc','0'),('locked','0'),
                ('layoutInCell','1'),('allowOverlap','1')]:
        anc.set(k,v)
    sp_=etree.SubElement(anc,Q('wp','simplePos')); sp_.set('x','0'); sp_.set('y','0')
    pH=etree.SubElement(anc,Q('wp','positionH')); pH.set('relativeFrom','page')
    etree.SubElement(pH,Q('wp','posOffset')).text=str(emu(x))
    pV=etree.SubElement(anc,Q('wp','positionV')); pV.set('relativeFrom','page')
    etree.SubElement(pV,Q('wp','posOffset')).text=str(emu(y))
    ex=etree.SubElement(anc,Q('wp','extent')); ex.set('cx',str(emu(w))); ex.set('cy',str(emu(h)))
    ee=etree.SubElement(anc,Q('wp','effectExtent'))
    for k in 'l','t','r','b': ee.set(k,'0')
    etree.SubElement(anc,Q('wp','wrapNone'))
    dp=etree.SubElement(anc,Q('wp','docPr')); dp.set('id',str(eid)); dp.set('name',f'S{eid}')
    etree.SubElement(anc,Q('wp','cNvGraphicFramePr'))
    g=etree.SubElement(anc,Q('a','graphic'))
    gd=etree.SubElement(g,Q('a','graphicData'))
    gd.set('uri','http://schemas.microsoft.com/office/word/2010/wordprocessingShape')
    wsp=etree.SubElement(gd,Q('wps','wsp'))
    cnv=etree.SubElement(wsp,Q('wps','cNvSpPr'))
    etree.SubElement(cnv,Q('a','spLocks')).set('noChangeArrowheads','1')
    spr=etree.SubElement(wsp,Q('wps','spPr'))
    xfm=etree.SubElement(spr,Q('a','xfrm'))
    if flipV: xfm.set('flipV','1')
    if flipH: xfm.set('flipH','1')
    etree.SubElement(xfm,Q('a','off')).set('x','0'); xfm[-1].set('y','0')
    etree.SubElement(xfm,Q('a','ext')).set('cx',str(emu(w))); xfm[-1].set('cy',str(emu(h)))
    pg=etree.SubElement(spr,Q('a','prstGeom')); pg.set('prst',prst)
    avl=etree.SubElement(pg,Q('a','avLst'))
    if adj:
        for nm,vl in adj:
            gd2=etree.SubElement(avl,Q('a','gd')); gd2.set('name',nm); gd2.set('fmla',f'val {vl}')
    if fill:
        sf=etree.SubElement(spr,Q('a','solidFill'))
        etree.SubElement(sf,Q('a','srgbClr')).set('val',fill)
    else:
        etree.SubElement(spr,Q('a','noFill'))
    ln=etree.SubElement(spr,Q('a','ln'))
    if border:
        ln.set('w',str(bw))
        bsf=etree.SubElement(ln,Q('a','solidFill'))
        etree.SubElement(bsf,Q('a','srgbClr')).set('val',border)
        if dashed: etree.SubElement(ln,Q('a','prstDash')).set('val','dash')
    else:
        etree.SubElement(ln,Q('a','noFill'))
    if line1 is not None:
        tx=etree.SubElement(wsp,Q('wps','txbx'))
        txc=etree.SubElement(tx,Q('w','txbxContent'))
        for txt,bold in [(line1,True)]+([(line2,False)] if line2 else []):
            if txt is None: continue
            p_=etree.SubElement(txc,Q('w','p'))
            pp=etree.SubElement(p_,Q('w','pPr'))
            jc=etree.SubElement(pp,Q('w','jc')); jc.set(Q('w','val'),'center')
            spc=etree.SubElement(pp,Q('w','spacing'))
            spc.set(Q('w','before'),'0'); spc.set(Q('w','after'),'0')
            r_=etree.SubElement(p_,Q('w','r'))
            rp=etree.SubElement(r_,Q('w','rPr'))
            if bold: etree.SubElement(rp,Q('w','b'))
            cl=etree.SubElement(rp,Q('w','color'))
            cl.set(Q('w','val'),tc if bold else '6B7280')
            sz=etree.SubElement(rp,Q('w','sz'))
            sz.set(Q('w','val'),str(fs*2 if bold else max(14,(fs-1)*2)))
            fn=etree.SubElement(rp,Q('w','rFonts'))
            fn.set(Q('w','ascii'),'Calibri'); fn.set(Q('w','hAnsi'),'Calibri')
            t_=etree.SubElement(r_,Q('w','t')); t_.text=txt
    bp=etree.SubElement(wsp,Q('wps','bodyPr'))
    bp.set('anchor','ctr'); bp.set('wrap','square')
    bp.set('lIns','68580'); bp.set('rIns','68580')
    bp.set('tIns','45720'); bp.set('bIns','45720')
    return mc

def add(para, x, y, w, h, prst, fill=None,
        border=None, bw=19050, dashed=False,
        line1=None, line2=None, fs=10, tc='1E293B',
        flipV=False, flipH=False, adj=None):
    r = OxmlElement('w:r')
    r.append(_shape(x,y,w,h,prst,fill,border,bw,dashed,
                    line1,line2,fs,tc,flipV,flipH,adj))
    para._p.append(r)

# ── Colours ───────────────────────────────────────────────────────────────────
PRC_F='FEF9C3'; PRC_B='D97706'   # process  – yellow
TRM_F='FECDD3'; TRM_B='BE123C'   # terminal – pink
DEC_F='FED7AA'; DEC_B='EA580C'   # decision – orange
IOP_F='BFDBFE'; IOP_B='1D4ED8'   # I-O      – blue
CST_F='EDE9FE'; CST_B='7C3AED'   # costing  – purple
DOC_F='D1FAE5'; DOC_B='047857'   # doc chip – mint
SEC_F='F0FDF4'; SEC_B='15803D'   # section  – pale green
FIN_F='BBF7D0'; FIN_B='064E3B'   # final
YES_F='DCFCE7'; YES_B='16A34A'
NO_F ='FEE2E2'; NO_B ='DC2626'
CON  ='475569'
HDR_F='F1F5F9'; HDR_B='CBD5E1'
BPT=19050; TPT=12700

# ── Document ──────────────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.page_width=Inches(8.5); s.page_height=Inches(11.0)
    s.top_margin=Inches(0.5);  s.bottom_margin=Inches(0.5)
    s.left_margin=Inches(0.5); s.right_margin=Inches(0.5)

p1 = doc.add_paragraph()
_pb = doc.add_paragraph()
_r  = _pb.add_run()
_br = OxmlElement('w:br'); _br.set(DQN('w:type'),'page'); _r._r.append(_br)
p2  = doc.add_paragraph()

# ── Layout ────────────────────────────────────────────────────────────────────
CX  = 4.25
BW  = 4.6;  BH  = 0.52
OW  = 1.6;  OH  = 0.48
DW  = 2.5;  DH  = 0.90
BX  = CX - BW/2
RR  = [('adj','8000')]

# Connector geometry
SW=0.022; HW=0.14; HH=0.10

def conn_down(para, y, h=0.28, color=CON):
    sh = max(0.01, h-HH)
    add(para, CX-SW/2, y,    SW, sh, 'rect',     color, border=None)
    add(para, CX-HW/2, y+sh, HW, HH, 'triangle', color, border=None, flipV=True)
    return y+h

def conn_right(para, lx, cy, length, color=CON):
    sh = max(0.01, length-HH)
    add(para, lx,    cy-SW/2, sh, SW, 'rect',     color, border=None)
    add(para, lx+sh, cy-HW/2, HH, HW, 'triangle', color, border=None,
        adj=[('adj','33333')])
    return lx+length

def conn_left_arrow(para, rx, cy, length, color=CON):
    """Arrow pointing left: arrowhead at left end."""
    sh = max(0.01, length-HH)
    # shaft
    add(para, rx-length+HH, cy-SW/2, sh, SW, 'rect', color, border=None)
    # arrowhead (rightArrow flipH = left arrow)
    add(para, rx-length, cy-HW/2, HH, HW, 'triangle', color, border=None,
        adj=[('adj','33333')], flipH=True)

def conn_up(para, x, y_top, height, color=CON):
    """Thin vertical bar (no arrowhead) used for loop sides."""
    add(para, x-SW/2, y_top, SW, height, 'rect', color, border=None)

# Shape helpers
def terminal(para, y, text):
    add(para, CX-OW/2, y, OW, OH, 'ellipse', TRM_F,
        border=TRM_B, bw=BPT, line1=text, fs=11, tc='881337')
    return y+OH

def process(para, y, text, sub=None, color_f=PRC_F, color_b=PRC_B):
    add(para, BX, y, BW, BH, 'roundRect', color_f,
        border=color_b, bw=BPT, line1=text, line2=sub, fs=10, tc='78350F', adj=RR)
    return y+BH

def io_box(para, y, text, sub=None):
    add(para, BX+0.2, y, BW-0.4, BH, 'parallelogram', IOP_F,
        border=IOP_B, bw=BPT, dashed=True, line1=text, line2=sub, fs=9, tc='1E3A8A')
    return y+BH

def decision(para, y, text='APPROVED?', fill=DEC_F, border=DEC_B, tc='7C2D12'):
    add(para, CX-DW/2, y, DW, DH, 'diamond', fill,
        border=border, bw=BPT, line1=text, fs=12, tc=tc)
    return y+DH, y+DH/2

def label_pill(para, x, cy, text, fill, br, tc):
    W2=0.40; H2=0.22
    add(para, x-W2/2, cy-H2/2, W2, H2, 'roundRect', fill,
        border=br, bw=TPT, line1=text, fs=8, tc=tc, adj=[('adj','25000')])

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1
# ══════════════════════════════════════════════════════════════════════════════

# Header bar
add(p1, 0.5, 0.42, 7.5, 0.86, 'roundRect', HDR_F,
    border=HDR_B, bw=TPT,
    line1='ED MODULE  —  BUSINESS WORKFLOW',
    line2='Buyer Order  ·  Price Review  ·  PI & LC  ·  Approval  ·  Production  ·  Export',
    fs=13, tc='0F172A', adj=[('adj','6000')])
add(p1, CX-2.3, 1.30, 4.6, 0.038, 'rect', '3B82F6', border=None)

y = 1.50

# START
y = terminal(p1, y, 'Start')
y = conn_down(p1, y)

# 1 – Buyer
y = process(p1, y, 'Buyer Places Order',
            'Items  ·  Quantity  ·  Delivery  ·  Commercial Requirement')
y = conn_down(p1, y)

# 2 – Marketing  ← remember position for loop return
mkt_top = y
y = process(p1, y, 'Marketing Receives & Reviews Order',
            'Understands requirements  ·  Sends to Costing for price review')
mkt_bot = y
mkt_mid = (mkt_top + mkt_bot) / 2
y = conn_down(p1, y)

# 3 – Costing / Price Review
y = process(p1, y, 'Costing  —  Price Review',
            'Calculates / reviews cost  ·  Approves or revises price',
            color_f=CST_F, color_b=CST_B)
y = conn_down(p1, y)

# 4 – Price Approved? decision
pa_top = y
pa_bot, pa_mid = decision(p1, y, 'PRICE\nAPPROVED?',
                          fill=DEC_F, border=DEC_B, tc='7C2D12')

# ── YES branch (price approved – continue down) ───────────────────────────────
label_pill(p1, CX, pa_bot+0.08, 'Yes', YES_F, YES_B, '166534')
conn_down(p1, pa_bot+0.30, 0.26, color=YES_B)

# ── NO branch (revise price – loop back to Marketing) ────────────────────────
NO_START_X = CX + DW/2
label_pill(p1, NO_START_X+0.28, pa_mid-0.28, 'No', NO_F, NO_B, '991B1B')

# Horizontal right from diamond
conn_right(p1, NO_START_X+0.02, pa_mid, 0.52, color=NO_B)

# "Revise Price" box
RV_X = NO_START_X + 0.56
RV_W = 1.85; RV_H = 0.56
RV_Y = pa_mid - RV_H/2
add(p1, RV_X, RV_Y, RV_W, RV_H, 'roundRect', NO_F,
    border=NO_B, bw=BPT,
    line1='Revise Price', line2='Return to Marketing',
    fs=9, tc='9F1239', adj=[('adj','10000')])

# Loop connector – right side going up then left to Marketing
LOOP_X = RV_X + RV_W + 0.06   # right edge of loop
# Vertical up segment (from mkt_mid to pa_mid, right side)
conn_up(p1, LOOP_X, mkt_mid, pa_mid - mkt_mid, color=NO_B)
# Horizontal left connector returning to Marketing right edge
MKT_RIGHT = BX + BW
ret_len    = LOOP_X - MKT_RIGHT
conn_left_arrow(p1, LOOP_X, mkt_mid, ret_len, color=NO_B)

# ── Continue YES flow ─────────────────────────────────────────────────────────
y = pa_bot + 0.58   # after YES conn

# 5 – Commercial enters SO
y = process(p1, y, 'Commercial Enters Sales Order Number')

# I-O auto-fill
io_y = y + 0.08
y    = io_box(p1, io_y,
              'System Auto-Fill  —  External Source',
              'Buyer · Items · Qty · Price · PI No · DC No · LC No · Dates')
y = conn_down(p1, y, 0.30)

# 6 – PI
y = process(p1, y, 'Commercial Creates / Checks Proforma Invoice (PI)',
            'Foundation document for all downstream forms')
y = conn_down(p1, y)

# 7 – LC
y = process(p1, y, 'Commercial Checks Letter of Credit (LC)',
            'Required before export & bank documents can be prepared')
y = conn_down(p1, y)

# 8 – Marketing Final Review
y = process(p1, y, 'Marketing Reviews PI & LC Details',
            'Approves  or  returns to Commercial for correction')

# Continued arrow (no arrowhead – connects to page 2)
add(p1, CX-SW/2, y+0.02, SW, 0.22, 'rect', CON, border=None)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2
# ══════════════════════════════════════════════════════════════════════════════
y2 = 0.50

# Entry arrow from page 1
conn_down(p2, y2, 0.30, color=CON)
y2 += 0.32

# 9 – APPROVED? (commercial approval)
ap_bot, ap_mid = decision(p2, y2)

label_pill(p2, CX, ap_bot+0.08, 'Yes', YES_F, YES_B, '166534')
conn_down(p2, ap_bot+0.30, 0.26, color=YES_B)

# NO rejection
label_pill(p2, CX+DW/2+0.28, ap_mid-0.28, 'No', NO_F, NO_B, '991B1B')
conn_right(p2, CX+DW/2+0.02, ap_mid, 0.52, color=NO_B)
add(p2, CX+DW/2+0.56, ap_mid-0.28, 1.95, 0.56, 'roundRect', NO_F,
    border=NO_B, bw=BPT,
    line1='Return to Commercial', line2='for Correction',
    fs=9, tc='9F1239', adj=[('adj','10000')])

y2 = ap_bot + 0.58

# 10 – Factory
y2 = process(p2, y2, 'Factory Starts Production',
             'Approved orders only  ·  Goods prepared  ·  Coordinates dispatch')
y2 = conn_down(p2, y2)

# ── Document section helper ───────────────────────────────────────────────────
def doc_section(para, y, title, subtitle, chips):
    cols=3; cw=(BW-0.10*(cols+1))/cols; ch=0.44; hh=0.46; vg=0.10
    rows=(len(chips)+cols-1)//cols; sec_h=hh+rows*(ch+vg)+vg
    add(para, BX, y, BW, sec_h, 'roundRect', SEC_F,
        border=SEC_B, bw=BPT,
        line1=title, line2=subtitle, fs=10, tc='065F46', adj=[('adj','6000')])
    for i,ch_ in enumerate(chips):
        row=i//cols; col=i%cols
        rn=min(cols,len(chips)-row*cols)
        rw=rn*cw+(rn-1)*0.10; sx=CX-rw/2
        cx_=sx+col*(cw+0.10)+cw/2; cy_=y+hh+vg+row*(ch+vg)
        add(para, cx_-cw/2, cy_, cw, ch, 'roundRect', DOC_F,
            border=DOC_B, bw=TPT, line1=ch_, fs=8, tc='065F46',
            adj=[('adj','12000')])
    return y+sec_h

# 11 – Documents Phase 1
y2 = doc_section(p2, y2, 'Documents  —  Phase 1', 'Delivery & Packing',
                 ['Commercial\nInvoice','Packing\nList','Delivery\nChallan',
                  'Truck\nChallan','Challan\nSheet'])
y2 = conn_down(p2, y2)

# 12 – Documents Phase 2
y2 = doc_section(p2, y2, 'Documents  —  Phase 2', 'Export & Bank',
                 ['Certificate\nof Origin','Beneficiary\nCertificate',
                  'Bill of\nExchange'])
y2 = conn_down(p2, y2)

# END terminal
EW2 = OW+0.6
add(p2, CX-EW2/2, y2, EW2, OH+0.14, 'ellipse', FIN_F,
    border=FIN_B, bw=BPT,
    line1='Forwarding Pack  —  End',
    line2='Final set ready for export & bank',
    fs=11, tc='064E3B')

# ── Legend ────────────────────────────────────────────────────────────────────
y2 += OH + 0.5; lx = 0.5
for lbl, f, b, tc_ in [
    ('Terminal',    TRM_F, TRM_B, '881337'),
    ('Process',     PRC_F, PRC_B, '78350F'),
    ('Costing',     CST_F, CST_B, '5B21B6'),
    ('Decision',    DEC_F, DEC_B, '7C2D12'),
    ('Data / I-O',  IOP_F, IOP_B, '1E3A8A'),
    ('Document',    DOC_F, DOC_B, '065F46'),
]:
    add(p2, lx+0.04, y2+0.04, 0.18, 0.18, 'roundRect', f,
        border=b, bw=TPT, adj=[('adj','15000')])
    add(p2, lx+0.26, y2, 1.05, 0.26, 'rect', 'FFFFFF', border=None,
        line1=lbl, fs=7, tc='334155')
    lx += 1.35

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   'workflow_flowchart.docx')
doc.save(OUT)
print('Saved:', OUT)
