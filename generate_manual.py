"""
Generate ATS User Manual as a .docx file
Run: python generate_manual.py
Output: ATS_User_Manual.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

SS = r'e:\xampp\htdocs\ed_module\screenshots'

def add_screenshot(filename, caption=''):
    path = os.path.join(SS, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(6.0))
    if caption:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(10)
        r = p2.add_run(caption)
        set_font(r, 9, italic=True, color=GREY)

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Colour palette ──────────────────────────────────────────────────────────
INDIGO  = RGBColor(0x63, 0x66, 0xF1)   # primary brand colour
DARK    = RGBColor(0x1E, 0x1E, 0x2E)
GREY    = RGBColor(0x64, 0x74, 0x8B)
GREEN   = RGBColor(0x16, 0xA3, 0x4A)
AMBER   = RGBColor(0xD9, 0x77, 0x06)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT   = RGBColor(0xF5, 0xF7, 0xFF)

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_font(run, 18, bold=True, color=INDIGO)
    elif level == 2:
        set_font(run, 13, bold=True, color=DARK)
    elif level == 3:
        set_font(run, 11, bold=True, color=GREY)
    return p

def add_body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, 11, color=DARK)
    return p

def add_step(number, title, description):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.0)
    r1 = p.add_run(f"Step {number}  ")
    set_font(r1, 11, bold=True, color=INDIGO)
    r2 = p.add_run(title)
    set_font(r2, 11, bold=True, color=DARK)
    if description:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent  = Inches(0.55)
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(6)
        run = p2.add_run(description)
        set_font(run, 10.5, color=GREY)

def add_bullet(text, sub=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.5 if not sub else 0.85)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, 10.5, color=DARK)

def add_note(text, kind='note'):
    """kind: 'note' | 'tip' | 'warning'"""
    labels = {'note': ('ℹ  Note', INDIGO), 'tip': ('✔  Tip', GREEN), 'warning': ('⚠  Important', AMBER)}
    label, color = labels.get(kind, labels['note'])
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(label + '  ')
    set_font(r1, 10, bold=True, color=color)
    r2 = p.add_run(text)
    set_font(r2, 10, italic=True, color=GREY)

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run('─' * 72)
    run.font.color.rgb = RGBColor(0xE0, 0xE3, 0xFF)
    run.font.size = Pt(7)

def add_page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run('Accessories Tracking System')
set_font(r, 28, bold=True, color=INDIGO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ATS — Order & Export Documentation')
set_font(r, 14, color=GREY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
r = p.add_run('User Manual')
set_font(r, 20, bold=True, color=DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
r = p.add_run('For non-technical users  ·  All roles')
set_font(r, 11, color=GREY, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run(f'Prepared: {datetime.date.today().strftime("%B %Y")}')
set_font(r, 10, color=GREY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Zaber & Zubair Accessories Ltd.')
set_font(r, 11, bold=True, color=DARK)

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual — no auto TOC field)
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Table of Contents', 1)

toc = [
    ('1.', 'Getting Started — Login'),
    ('2.', 'The Screen Layout'),
    ('3.', 'Master Data'),
    ('   3.1', 'Customer List'),
    ('   3.2', 'Create a New Customer Profile'),
    ('   3.3', 'Item Master'),
    ('   3.4', 'User Management (Admin only)'),
    ('4.', 'Order & Documents Workflow — Overview'),
    ('5.', 'The Order ID Bar — Loading & Creating Orders'),
    ('6.', 'Step-by-Step: Order & Documents'),
    ('   Step 1', 'Marketing Intake'),
    ('   Step 2', 'Costing Review'),
    ('   Step 3', 'PI (Proforma Invoice / Sales)'),
    ('   Step 4', 'Marketing Approval'),
    ('   Step 5', 'LC (Letter of Credit)'),
    ('   Step 6', 'Bill of Exchange'),
    ('   Step 7', 'Commercial Invoice'),
    ('   Step 8', 'Packing List'),
    ('   Step 9', 'Delivery Challan'),
    ('   Step 10', 'Truck Challan'),
    ('   Step 11', 'Certificate of Origin'),
    ('   Step 12', "Beneficiary's Certificate"),
    ('   Step 13', 'Forwarding'),
    ('   Step 14', 'Challan Sheet'),
    ('7.', 'Saving vs. Moving to the Next Step'),
    ('8.', 'Signing Out'),
    ('9.', 'Frequently Asked Questions'),
]

for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'{num:<10}')
    set_font(r1, 10.5, bold=('.' not in num.strip() or num.strip().endswith('.')), color=INDIGO if '.' not in num.strip() else DARK)
    r2 = p.add_run(title)
    set_font(r2, 10.5, color=DARK)

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. LOGIN
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1.  Getting Started — Login', 1)

add_body(
    'The Accessories Tracking System (ATS) runs in your web browser. No software needs to be installed on your computer. '
    'Your IT team will give you the web address (URL).'
)

add_heading('How to Log In', 2)
add_bullet('Open your web browser (Chrome, Edge, or Firefox are recommended).')
add_bullet('Type the ATS address in the address bar and press Enter.')
add_bullet('You will see the login page with the ATS logo.')
add_bullet('Enter your Email address in the first box.')
add_bullet('Enter your Password in the second box.')
add_bullet('Click the blue Sign In button.')
add_bullet('You will be taken to the Dashboard automatically.')

add_screenshot('01_login.png', 'Figure 1 — ATS Login Page')
add_note('If you see "Invalid email or password", check that Caps Lock is off and try again. Contact your administrator if you are still unable to log in.', 'warning')
add_note('Your session stays active as long as the browser tab is open. Closing the tab will require you to log in again.', 'note')

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  2. SCREEN LAYOUT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2.  The Screen Layout', 1)

add_body(
    'Every page in ATS has the same structure. Once you learn where things are, you can move around quickly.'
)

add_screenshot('02_dashboard.png', 'Figure 2 — Dashboard overview')
add_heading('Top Navigation Bar', 2)
add_body('The dark bar at the very top contains three main buttons:')
add_bullet('Master Data — takes you to the customer and item lists.')
add_bullet('Order & Documents — takes you to the full order workflow (this is where most daily work happens).')
add_bullet('Dashboard — gives a summary view of all orders.')
add_bullet('Your name and role are shown on the right. Next to them is the Sign Out link.')

add_heading('Tab Row (second row)', 2)
add_body(
    'Below the top bar is a row of clickable tabs separated by arrows (→). '
    'These tabs show every step in the current section. The currently active tab is highlighted in purple. '
    'You can click any tab to jump directly to that page.'
)

add_heading('Work Order Bar (Order & Documents only)', 2)
add_body(
    'When you are in the Order & Documents section, a dark bar appears below the tabs. '
    'This bar shows which order is currently loaded (e.g. ORD-2026-0012). '
    'All the pages you fill in will be saved against this order.'
)

add_heading('Main Form Area', 2)
add_body('The large white area below is where you read and fill in information for the current step.')

add_heading('Action Buttons (bottom of the page)', 2)
add_body('Every page has buttons at the bottom:')
add_bullet('Previous — goes back to the previous step.')
add_bullet('Next … → — saves the current page and moves to the next step.')
add_bullet('Save Intake / Save PI etc. — saves the current page without moving forward (available on some steps).')

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  3. MASTER DATA
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3.  Master Data', 1)

add_body(
    'Master Data is a one-time setup area. You add customers and items here before you can create orders for them. '
    'Click Master Data in the top bar to enter this section.'
)

add_screenshot('03_customers.png', 'Figure 3 — Customer List')
# 3.1
add_heading('3.1  Customer List', 2)
add_body(
    'This page shows every customer that has been entered into the system. '
    'You can scroll through the list or use the search box at the top to find a customer by name.'
)
add_bullet('Each row shows the company name, type, chairman name, and current approval stage.')
add_bullet('Click a customer row to view their full profile.')
add_note('Customers must go through an approval process (Sales Person → Commercial → Team Leader → etc.) before they appear in the order dropdown.', 'note')

# 3.2
add_heading('3.2  Create a New Customer Profile', 2)
add_body(
    'Click the Create Profile tab to open the new customer form. '
    'Fill in every section as completely as possible. Fields marked with a red * are required.'
)

sections_cp = [
    ('1. Customer Information', 'Company name, type, head office address, factory address, chairman/director details.'),
    ('2. Business & Compliance', 'Trade license, BIN, TIN, bond license, bank name, factory certifications (BSCI, WRAP, etc.).'),
    ('3. Production Capability', 'Factory type, monthly capacity, number of machines, major buyers and products.'),
    ('4. Commercial Assessment', 'Expected monthly business, LC terms, BBLC terms, payment currency, delivery terms.'),
    ('5. Product Interest', 'Tick all product types this customer is likely to order (Carton, Poly, Label, etc.).'),
    ('6. Competitor Analysis', 'Current supplier, their price, strengths, weaknesses, and reason for switching to you.'),
    ('7. Risk Assessment', 'Financial risk level, payment history, and recommended credit limit.'),
    ('8. Price Approval Matrix', 'List products with their existing, target, and approved prices.'),
    ('9. Document Checklist', 'Tick which compliance documents have been received from the customer.'),
    ('Signature', 'Upload a photo or scan of the Sales Person\'s signature before submitting.'),
]
for title, desc in sections_cp:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'{title}  ')
    set_font(r1, 10.5, bold=True, color=INDIGO)
    r2 = p.add_run(desc)
    set_font(r2, 10.5, color=DARK)

add_body('')
add_bullet('Click Save Draft at any time to save your progress without submitting. You can come back later.')
add_bullet('When everything is complete, click Submit Profile. A green confirmation message will appear and you will be taken back to the Customer List.')
add_screenshot('04_create_cust.png', 'Figure 4 — Create Customer Profile form')
add_note('Submitting a profile sends it to the first approval stage. The approver must then log in and approve it before the customer appears in the order dropdown.', 'note')

# 3.3
add_heading('3.3  Item Master', 2)
add_body(
    'This page lists all products and accessories the company sells. '
    'Each item shows its Product Line, Item Name, Grade (for Carton items), Paper Combination, and Base Price.'
)
add_bullet('To add a new item, click the + Add Item button.')
add_bullet('Select a Product Line first — the item name list will update automatically.')
add_bullet('For Carton items, a Grade dropdown and Paper Combination dropdown will appear.')
add_bullet('Enter the Base Price and click Save Item.')
add_note('Items in this master list are available for selection when filling in a Marketing Intake order.', 'note')

# 3.4
add_screenshot('05_item_master.png', 'Figure 5 — Item Master')
add_heading('3.4  User Management (Admin only)', 2)
add_body(
    'Only users with the Admin role can see the Users tab. '
    'Here you can create new user accounts, change roles, reset passwords, and deactivate accounts.'
)

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. WORKFLOW OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4.  Order & Documents Workflow — Overview', 1)

add_body(
    'Every export order follows the same 14-step journey from first contact with a customer to the final bank documents. '
    'The tabs at the top show all steps in order. You must complete them from left to right — '
    'each step builds on information from the previous ones.'
)

steps_overview = [
    ('1',  'Marketing Intake',          'Record the customer order — PO numbers, items, quantities, prices.'),
    ('2',  'Costing Review',            'Internal review of item prices; revise if needed.'),
    ('3',  'PI',                        'Create the Proforma Invoice (PI) to send to the customer.'),
    ('4',  'Marketing Approval',        'Marketing team signs off on the order details.'),
    ('5',  'LC',                        'Enter the Letter of Credit details received from the bank.'),
    ('6',  'Bill of Exchange',          'Fill in the Bill of Exchange (bank negotiation document).'),
    ('7',  'Commercial Invoice',        'Generate the Commercial Invoice for the shipment.'),
    ('8',  'Packing List',              'Auto-generated packing details based on the Commercial Invoice.'),
    ('9',  'Delivery Challan',          'Record the delivery details for the goods.'),
    ('10', 'Truck Challan',             'Record truck/transport details.'),
    ('11', 'Certificate of Origin',     'Auto-generated certificate stating goods are from Bangladesh.'),
    ('12', "Beneficiary's Certificate", 'Certificate confirming accessories are for export garments.'),
    ('13', 'Forwarding',                'Cover letter submitted to the bank with all negotiation documents.'),
    ('14', 'Challan Sheet',             'Quality control and delivery inspection record.'),
]

for num, name, desc in steps_overview:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f'  {num:>2}.  ')
    set_font(r1, 10.5, bold=True, color=INDIGO)
    r2 = p.add_run(f'{name:<32}')
    set_font(r2, 10.5, bold=True, color=DARK)
    r3 = p.add_run(desc)
    set_font(r3, 10, color=GREY)

add_note(
    'Steps 8–12 (Packing List, Delivery Challan, Truck Challan, Certificate of Origin, Beneficiary\'s Certificate) '
    'are automatically filled from data you entered in earlier steps. You only need to review them and click Next.',
    'tip'
)

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. ORDER ID BAR
# ══════════════════════════════════════════════════════════════════════════════
add_heading('5.  The Order ID Bar — Loading & Creating Orders', 1)

add_body(
    'The dark bar just below the tabs is called the Order ID Bar. '
    'It tells you which order you are currently working on. '
    'Before filling in any step, you must either load an existing order or create a new one.'
)

add_heading('Creating a New Order', 2)
add_bullet('Click the + New Order button on the right side of the Order ID Bar.')
add_bullet('A new Order ID will be generated automatically (e.g. ORD-2026-0013) and shown in the bar.')
add_bullet('You are now ready to fill in the Marketing Intake (Step 1).')
add_note('Creating a new order does not ask for any information yet — it just reserves an ID. All details go in the Marketing Intake form.', 'note')

add_heading('Loading an Existing Order', 2)
add_bullet('Type the Order ID (e.g. ORD-2026-0012) into the search box in the Order ID Bar.')
add_bullet('Press Enter or click the Load Order button.')
add_bullet('The bar will update to show the customer name, current step, and date.')
add_bullet('All previously saved data for this order will load automatically on each page.')
add_note('The system remembers the last order you loaded in this browser tab. If you close and reopen the tab, you may need to re-load the order.', 'note')

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. STEP-BY-STEP
# ══════════════════════════════════════════════════════════════════════════════
add_heading('6.  Step-by-Step: Order & Documents', 1)

# ── Step 1 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 1  —  Marketing Intake', 2)
add_body(
    'This is the starting point for every order. The Sales Person fills in the full details '
    'of what the customer has ordered.'
)
add_heading('What to fill in:', 3)
add_bullet('Customer — select from the dropdown (must be a customer already in the system).')
add_bullet('Sales Person — fills automatically once a customer is selected.')
add_bullet('Intake Date — today\'s date fills automatically; change if needed.')
add_bullet('Sub / Work Order Description — a short description of what is being ordered.')
add_bullet('Paper Quality — e.g. "350 GSM".')
add_heading('For each PO (Purchase Order):', 3)
add_bullet('PO Number — the customer\'s purchase order number.')
add_bullet('End Buyer — the brand/buyer the customer is making goods for (e.g. H&M).')
add_bullet('TRIMS / IPO No. — TRIMS or internal PO reference if applicable.')
add_bullet('Design, Order No., Type, Delivery Date — fill as per customer PO.')
add_bullet('Item Lines — add each item with Product Line, Item Name, Art/Size, Grade, Paper Combination, Qty, Unit, and Unit Price.')
add_bullet('Click + Add Row to add more items to a PO.')
add_bullet('Click + Add Another PO if the customer has more than one PO in this order.')
add_heading('Saving & Sending to Costing:', 3)
add_bullet('Click 💾 Save Intake to save without moving forward (safe to do anytime).')
add_bullet('Click Send To Costing → when the intake is complete. You will be asked to confirm and then taken to Step 2.')
add_screenshot('06_intake.png', 'Figure 6 — Step 1: Marketing Intake')
add_note('You can come back to edit the intake at any time — the system will show a yellow notice if the order has already moved to a later step, but you can still re-save.', 'tip')

# ── Step 2 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 2  —  Costing Review', 2)
add_body(
    'The Costing team reviews the item prices entered at intake. '
    'If a price needs to be changed, the revised price is entered here.'
)
add_bullet('All items from the intake are displayed in a table.')
add_bullet('The "Revised Price" column is editable. Enter a new price if the standard price needs adjustment.')
add_bullet('The status badge next to each row shows "Pending" (no revision) or "Revised" (price changed).')
add_bullet('Click Next: PI → when done. The revised prices will automatically carry into the PI.')
add_screenshot('07_costing.png', 'Figure 7 — Step 2: Costing Review')
add_note('If no prices need changing, simply click Next: PI → without entering anything.', 'tip')

# ── Step 3 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 3  —  PI (Proforma Invoice)', 2)
add_body(
    'The PI (Proforma Invoice) is the formal price quotation sent to the customer. '
    'Many fields carry over automatically from the intake; you mainly fill in the bank and LC details.'
)
add_bullet('Sales Order No. / Customer PO / Buyer Name — loaded from intake.')
add_bullet('PI Number — the internal PI number (e.g. LIZ-LO-CTN-26020001).')
add_bullet('PI Date — the date of the PI.')
add_bullet('Advising Bank — the bank that will advise the Letter of Credit.')
add_bullet('Consignee Address — the address of the buyer/customer.')
add_bullet('Consignee\'s Bank — the customer\'s bank details.')
add_bullet('Payment Terms, Delivery Terms, Validity, etc. — fill as per agreement.')
add_bullet('Item table — pre-filled from intake; adjust if needed.')
add_bullet('Click Next: Marketing → when done.')
add_screenshot('08_pi.png', 'Figure 8 — Step 3: Proforma Invoice (PI)')
add_note('The bank fields (Advising Bank, Consignee Bank) entered here will be used by all later documents automatically.', 'tip')

# ── Step 4 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 4  —  Marketing Approval', 2)
add_body(
    'The Marketing team reviews the PI details and records their approval.'
)
add_bullet('All order details are displayed for review.')
add_bullet('The Marketing person can add notes or comments.')
add_screenshot('09_marketing.png', 'Figure 9 — Step 4: Marketing Approval')
add_bullet('Click Next: LC → to proceed after approval.')

# ── Step 5 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 5  —  LC (Letter of Credit)', 2)
add_body(
    'Once the customer opens an LC at their bank, you enter those details here. '
    'This is one of the most important steps — the LC details flow into many later documents.'
)
add_bullet('LC Number — the LC number as shown on the bank document.')
add_bullet('LC Date — the date printed on the LC.')
add_bullet('LC Amount (USD) — the total value of the LC.')
add_bullet('Expiry Date — when the LC expires.')
add_bullet('Latest Shipment Date — the last date goods can be shipped under this LC.')
add_bullet('Concerning Receiver Bank — the bank receiving the LC on your side.')
add_bullet('LC Check Status — select the current LC status (Received, Pending, etc.).')
add_bullet('Click Next: Bill of Exchange → when done.')
add_screenshot('10_lc.png', 'Figure 10 — Step 5: Letter of Credit (LC)')
add_note('Double-check the LC Number and LC Date carefully — these appear on all bank documents and must match the original LC exactly.', 'warning')

# ── Step 6 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 6  —  Bill of Exchange', 2)
add_body(
    'The Bill of Exchange (BOE) is a formal bank instruction to pay. '
    'Most fields are auto-filled from earlier steps.'
)
add_bullet('Source details (Sales Order, Customer PO, Buyer Name) — auto-filled from intake.')
add_bullet('Master LC No. / Date — auto-filled from the LC step.')
add_bullet('Pay To (Bank Name / Address) — auto-filled from the PI advising bank; edit if needed.')
add_bullet('Applicant details (name, factory address, IRC, TIN, VAT/BIN numbers) — fill if not already present.')
add_bullet('Exchange Amount (USD) — the amount to be paid.')
add_bullet('Amount in Words — write the amount in words (e.g. "Eight Thousand One Hundred Twenty Three and Cents Ninety Eight Only").')
add_bullet('Tenor — payment terms (e.g. "At Sight").')
add_bullet('Export Sales Contract No. / Date — if applicable.')
add_bullet('HS Code — the harmonised system code for the goods.')
add_bullet('Click Next: Commercial Invoice → when done.')

add_screenshot('11_exchange.png', 'Figure 11 — Step 6: Bill of Exchange')
# ── Step 7 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 7  —  Commercial Invoice', 2)
add_body(
    'The Commercial Invoice is the main shipment invoice submitted to the bank. '
    'Most details carry over automatically.'
)
add_bullet('Invoice No. / Date — the invoice number and date.')
add_bullet('Proforma No. / Date — auto-filled from PI.')
add_bullet('LC No. / Date — auto-filled from LC step.')
add_bullet('Beneficiary Name / Address — your company details.')
add_bullet('Consignee Name / Address — auto-filled from PI (buyer address).')
add_bullet('Advising Bank — auto-filled from PI.')
add_bullet('Consignee\'s Bank — auto-filled from PI.')
add_bullet('Item table — review quantities and amounts. Adjust if the final shipment differs from the PI.')
add_bullet('Total Amount (USD) — calculated automatically.')
add_bullet('Click Next: Packing List → when done.')

add_screenshot('12_commercial.png', 'Figure 12 — Step 7: Commercial Invoice')
# ── Step 8 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 8  —  Packing List', 2)
add_body(
    'The Packing List is generated automatically from the Commercial Invoice. '
    'Review the details and proceed.'
)
add_bullet('All items, quantities, and totals are pulled from the Commercial Invoice.')
add_bullet('Company header and LC details are filled from earlier steps.')
add_bullet('No data entry is required on this page in normal circumstances.')
add_bullet('Click Next: Delivery Challan → to proceed.')

add_screenshot('13_packing.png', 'Figure 13 — Step 8: Packing List')
# ── Step 9 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 9  —  Delivery Challan', 2)
add_body('The Delivery Challan records the physical delivery of goods. Review auto-filled details and proceed.')
add_bullet('Delivery date, LC details, and item lines are populated from earlier steps.')
add_bullet('Click Next: Truck Challan → to proceed.')

add_screenshot('14_delivery.png', 'Figure 14 — Step 9: Delivery Challan')
# ── Step 10 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 10  —  Truck Challan', 2)
add_body('The Truck Challan records the truck/transport used for delivery.')
add_bullet('Vehicle number, driver details, and route are filled here.')
add_bullet('Item and LC details are auto-filled.')
add_bullet('Click Next: Certificate of Origin → to proceed.')

add_screenshot('15_truck.png', 'Figure 15 — Step 10: Truck Challan')
# ── Step 11 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 11  —  Certificate of Origin', 2)
add_body(
    'This certificate is required for export and confirms the goods are of Bangladesh origin. '
    'It is auto-generated from your order data.'
)
add_bullet('The statement text is built automatically using the LC number, date, PI number, and customer name.')
add_bullet('Review the text and make sure all details are correct.')
add_bullet('Click Next: Beneficiary\'s Certificate → to proceed.')

add_screenshot('16_origin.png', 'Figure 16 — Step 11: Certificate of Origin')
# ── Step 12 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading("Step 12  —  Beneficiary's Certificate", 2)
add_body(
    "The Beneficiary's Certificate is submitted with the bank documents. "
    "It is auto-generated and confirms that the accessories were supplied for export garments."
)
add_bullet('Both statement paragraphs are built automatically from order data.')
add_bullet('Review the content — especially quantity, amount, and LC number.')
add_bullet("Click Next: Forwarding → to proceed.")

add_screenshot('17_beneficiary.png', "Figure 17 — Step 12: Beneficiary's Certificate")
# ── Step 13 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 13  —  Forwarding', 2)
add_body(
    'The Forwarding letter is the cover letter submitted to your bank along with all the negotiation documents.'
)
add_bullet('Date, Reference No., Bank Name, and LC details are auto-filled.')
add_bullet('The document list table shows the number of copies of each document to submit (e.g. Bill of Exchange: 2 copies).')
add_bullet('You can edit the copy counts in the table if the bank requires a different number.')
add_bullet('Amount in words is auto-filled from the Bill of Exchange.')
add_bullet('Click Next: Challan Sheet → to proceed.')

add_screenshot('18_forwarding.png', 'Figure 18 — Step 13: Forwarding letter')
# ── Step 14 ──────────────────────────────────────────────────────────────────
add_divider()
add_heading('Step 14  —  Challan Sheet', 2)
add_body(
    'The Challan Sheet is a Quality Assurance record. '
    'It is searched by PI number and shows the delivery and inspection details.'
)
add_bullet('Enter the PI number in the search box at the top and click Search PI.')
add_bullet('The challan table will populate with the delivery rows.')
add_bullet('The QA team fills in the Pass / Fail result and signs.')
add_bullet('Click Finish Pack when the order is fully complete.')
add_screenshot('19_challan.png', 'Figure 19 — Step 14: Challan Sheet')
add_note('This is the final step. Once Finish Pack is clicked, the order documentation is complete.', 'tip')

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  7. SAVING
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7.  Saving vs. Moving to the Next Step', 1)

add_body(
    'Understanding when data is saved helps you avoid losing work.'
)

rows = [
    ('Action',                   'What happens'),
    ('💾 Save Intake',           'Saves the current intake form. You stay on the same page. Use this often!'),
    ('Send To Costing →',        'Saves the intake AND advances the order to Step 2. Use when intake is complete.'),
    ('Next … →',                 'On Steps 2–14: saves the current page AND navigates to the next step.'),
    ('Previous',                 'Goes back to the previous step WITHOUT saving. Save first if needed.'),
    ('Closing the browser tab',  'Any unsaved changes on the current page will be lost.'),
]

table = doc.add_table(rows=len(rows), cols=2)
table.style = 'Table Grid'
for i, (action, result) in enumerate(rows):
    row = table.rows[i]
    c0 = row.cells[0].paragraphs[0]
    c1 = row.cells[1].paragraphs[0]
    r0 = c0.add_run(action)
    r1 = c1.add_run(result)
    if i == 0:
        set_font(r0, 10, bold=True, color=WHITE)
        set_font(r1, 10, bold=True, color=WHITE)
        row.cells[0]._element.get_or_add_tcPr()
        row.cells[1]._element.get_or_add_tcPr()
        for cell in row.cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '6366F1')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:val'), 'clear')
            cell._element.tcPr.append(shading)
    else:
        set_font(r0, 10, bold=True, color=INDIGO)
        set_font(r1, 10, color=DARK)

doc.add_paragraph()

add_note(
    'The Status indicator in the top summary bar (shows "Saved ✓" in green or "Draft" in amber) '
    'tells you whether the current page has been saved.',
    'tip'
)

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  8. SIGN OUT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('8.  Signing Out', 1)

add_body(
    'Always sign out when you have finished using the system, especially on shared computers.'
)
add_bullet('Click Sign Out in the top-right corner of any page.')
add_bullet('You will be returned to the login page immediately.')
add_bullet('Your session data is cleared — no one else can access your account from the same browser.')
add_note('Simply closing the browser window does NOT sign you out. Always use the Sign Out link.', 'warning')

add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  9. FAQ
# ══════════════════════════════════════════════════════════════════════════════
add_heading('9.  Frequently Asked Questions', 1)

faqs = [
    (
        'I filled in a page and clicked Next, but when I came back the data was gone. What happened?',
        'Make sure an order is loaded in the Order ID Bar before filling in any page. '
        'If "No order loaded" is shown, click + New Order or type the Order ID and click Load Order first.'
    ),
    (
        'I cannot find the customer in the Marketing Intake dropdown.',
        'The customer must be created in Master Data → Create Profile first, and then approved through '
        'all approval stages. Ask the approver to complete their step.'
    ),
    (
        'The page shows a yellow warning banner at the top.',
        'This means the order has already moved past this step. You can still edit and re-save the data here, '
        'but the order progress will not go backwards.'
    ),
    (
        'I made a mistake in the PI and the order has already moved to LC. Can I fix it?',
        'Yes. Click the PI tab in the navigation, make your corrections, and click Next: Marketing → again. '
        'The corrected data will flow into the later documents when they are viewed.'
    ),
    (
        'The "Send To Costing" button shows an error or nothing happens.',
        'Make sure (1) a Customer is selected, (2) an Intake Date is entered, and (3) an Order ID is loaded '
        'in the Order ID Bar. All three are required before saving.'
    ),
    (
        'Several pages (Packing List, Certificate of Origin, etc.) show dashes "—" instead of data.',
        'These pages are filled automatically from earlier steps. Load the order in the Order ID Bar '
        '(enter the Order ID and click Load Order) and the data will appear.'
    ),
    (
        'I need to add a new item that is not in the Item Master.',
        'Go to Master Data → Item Master and click + Add Item. After saving, the item will be available '
        'in the Marketing Intake item dropdown.'
    ),
    (
        'Can two people work on the same order at the same time?',
        'It is not recommended. If two people save the same page at the same time, the last save will overwrite '
        'the earlier one. Coordinate with your team to work on different steps simultaneously.'
    ),
]

for q, a in faqs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run('Q:  ' + q)
    set_font(r, 10.5, bold=True, color=DARK)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent  = Inches(0.25)
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after  = Pt(4)
    r2 = p2.add_run('A:  ' + a)
    set_font(r2, 10.5, color=GREY)

# ── Footer note ──────────────────────────────────────────────────────────────
add_divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
r = p.add_run('Accessories Tracking System  ·  Internal User Manual  ·  Zaber & Zubair Accessories Ltd.')
set_font(r, 9, color=GREY, italic=True)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r'e:\xampp\htdocs\ed_module\ATS_User_Manual.docx'
doc.save(out)
print(f'Saved: {out}')
