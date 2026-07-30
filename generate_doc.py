"""
ATS Role Documentation Generator — 4 separate DOCX files, one per role.
Run: python generate_doc.py
Output: ATS_Marketing.docx, ATS_Costing.docx, ATS_Production.docx, ATS_Commercial.docx
"""

import asyncio, sys, io
from pathlib import Path
from datetime import date

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

BASE_URL   = "http://localhost/ed_module"
DEMO_ORDER = "ORD-2026-0019"
OUT_DIR    = Path("doc_screenshots")
VIEWPORT   = {"width": 1440, "height": 900}

ROLES = [
    {
        "id":       "marketing",
        "name":     "Marketing Officer",
        "email":    "marketing@zzal.com",
        "password": "mkt@123",
        "color":    (236, 72, 153),
        "docx":     "ATS_Marketing.docx",
        "tagline":  "Manages the first impression — from buyer PO intake to sales summary.",
        "pages": [
            {
                "id":    "marketing-intake",
                "url":   "/pages/marketing-intake.php",
                "title": "Marketing Intake",
                "desc":  "Record customer PO details, buyer info, product lines, item specs, quantities and unit prices. Each product line shows relevant spec fields automatically.",
            },
            {
                "id":    "marketing",
                "url":   "/pages/marketing.php",
                "title": "Marketing Summary",
                "desc":  "Review the confirmed sales snapshot with PO-wise item breakdown and totals. Collapsible PO blocks let you drill into each buyer's items.",
            },
        ],
    },
    {
        "id":       "costing",
        "name":     "Costing Officer",
        "email":    "costing@zzal.com",
        "password": "cost@123",
        "color":    (245, 158, 11),
        "docx":     "ATS_Costing.docx",
        "tagline":  "Reviews cost per item and approves the order before it enters production.",
        "pages": [
            {
                "id":    "costing-review",
                "url":   "/pages/costing-review.php",
                "title": "Costing Review",
                "desc":  "Check quantities, specs, and unit prices pulled from Marketing Intake. Revised prices can be entered per line. Click 'Approve & Send to Production' when satisfied.",
            },
        ],
    },
    {
        "id":       "production",
        "name":     "Production Officer",
        "email":    "production@zzal.com",
        "password": "prod@123",
        "color":    (16, 185, 129),
        "docx":     "ATS_Production.docx",
        "tagline":  "Confirms production has started so the Commercial team can raise the PI.",
        "pages": [
            {
                "id":    "production",
                "url":   "/pages/production.php",
                "title": "Production Entry",
                "desc":  "Mark the order as 'Started' once production begins. Set the start date and expected completion date. The PI step is unlocked for the Commercial team only after this is confirmed.",
            },
        ],
    },
    {
        "id":       "commercial_dept",
        "name":     "Commercial Officer",
        "email":    "commercial@zzal.com",
        "password": "comm@123",
        "color":    (99, 102, 241),
        "docx":     "ATS_Commercial.docx",
        "tagline":  "Handles all export documentation from PI through to final forwarding.",
        "pages": [
            {
                "id":    "sales",
                "url":   "/pages/sales.php",
                "title": "PI — Proforma Invoice",
                "desc":  "Search the ERP for the matching sales order, import PO line items, and generate the Proforma Invoice. Mark as Master PI when confirmed.",
            },
            {
                "id":    "lc",
                "url":   "/pages/lc.php",
                "title": "Letter of Credit",
                "desc":  "Record the LC details received from the buyer's bank — LC number, date, payment terms, amount, advising bank, and consignee bank.",
            },
            {
                "id":    "exchange",
                "url":   "/pages/exchange.php",
                "title": "Bill of Exchange",
                "desc":  "Enter applicant codes (IRC, TIN, VAT/BIN, Bank BIN, Bond License), HS Code, export sales contract, and carrier info. These auto-fill all downstream documents.",
            },
            {
                "id":    "commercial",
                "url":   "/pages/commercial.php",
                "title": "Commercial Invoice",
                "desc":  "Auto-populated invoice with beneficiary, consignee, LC details, and the full item table with quantities, unit prices, and totals.",
            },
            {
                "id":    "packing",
                "url":   "/pages/packing.php",
                "title": "Packing List",
                "desc":  "Auto-filled packing list showing all items, quantities, and ply. Includes signature slots for the authorized representative and consignee.",
            },
            {
                "id":    "delivery",
                "url":   "/pages/delivery.php",
                "title": "Delivery Challan",
                "desc":  "Dispatch document auto-filled from PI items, LC, exchange, and commercial data with all reference numbers.",
            },
            {
                "id":    "truck",
                "url":   "/pages/truck.php",
                "title": "Truck Challan",
                "desc":  "Transport document for truck dispatch with buyer name, item list, and all LC/contract references.",
            },
            {
                "id":    "origin",
                "url":   "/pages/origin.php",
                "title": "Certificate of Origin",
                "desc":  "Records country of origin, manufacturer details and the certifying signature.",
            },
            {
                "id":    "beneficiary",
                "url":   "/pages/beneficiary.php",
                "title": "Beneficiary Certificate",
                "desc":  "Confirms that goods match the LC terms — auto-populated with consignee and LC details.",
            },
            {
                "id":    "forwarding",
                "url":   "/pages/forwarding.php",
                "title": "Forwarding",
                "desc":  "Final freight and forwarding document capturing shipping details and freight forwarder information.",
            },
        ],
    },
]


# ── Screenshot capture ────────────────────────────────────────────────────────
async def capture_role(pw, role: dict) -> dict:
    browser = await pw.chromium.launch(headless=True)
    ctx     = await browser.new_context(viewport=VIEWPORT)
    page    = await ctx.new_page()
    shots   = {}
    rdir    = OUT_DIR / role["id"]
    rdir.mkdir(parents=True, exist_ok=True)

    await page.goto(BASE_URL + "/pages/login.php", wait_until="networkidle")
    await page.fill('input[name="email"]',    role["email"])
    await page.fill('input[name="password"]', role["password"])
    async with page.expect_navigation(timeout=12000):
        await page.click('button[type="submit"]')
    await page.wait_for_timeout(800)
    print(f"    Logged in: {role['email']}")

    for pg in role["pages"]:
        await page.goto(BASE_URL + pg["url"], wait_until="domcontentloaded")
        await page.wait_for_timeout(500)
        await page.evaluate(f"sessionStorage.setItem('ats_current_order_id', '{DEMO_ORDER}');")
        await page.reload(wait_until="domcontentloaded")
        try:
            await page.wait_for_load_state("networkidle", timeout=7000)
        except Exception:
            pass
        await page.wait_for_timeout(1400)
        await page.evaluate("window.scrollTo(0, 0)")
        await page.wait_for_timeout(400)
        path = str(rdir / f"{pg['id']}.png")
        await page.screenshot(path=path, full_page=True)
        shots[pg["id"]] = path
        print(f"    [ok] {pg['id']}")

    await browser.close()
    return shots


async def capture_all():
    from playwright.async_api import async_playwright
    OUT_DIR.mkdir(exist_ok=True)
    result = {}
    async with async_playwright() as pw:
        for role in ROLES:
            print(f"\n  [{role['name']}]")
            result[role["id"]] = await capture_role(pw, role)
    return result


# ── DOCX builder (one per role) ───────────────────────────────────────────────
def build_role_docx(role: dict, shots: dict):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    C = role["color"]

    def p(text="", size=11, bold=False, italic=False, color=(50,60,80),
          align=WD_ALIGN_PARAGRAPH.LEFT, after=6, before=0):
        para = doc.add_paragraph()
        para.alignment = align
        para.paragraph_format.space_after  = Pt(after)
        para.paragraph_format.space_before = Pt(before)
        if text:
            r = para.add_run(text)
            r.font.size = Pt(size); r.font.bold = bold
            r.font.italic = italic; r.font.color.rgb = RGBColor(*color)
        return para

    def h(text, size=14, color=(30,30,46), after=6, before=10):
        para = doc.add_paragraph()
        para.paragraph_format.space_after  = Pt(after)
        para.paragraph_format.space_before = Pt(before)
        r = para.add_run(text)
        r.font.size = Pt(size); r.font.bold = True
        r.font.color.rgb = RGBColor(*color)
        return para

    def divider():
        para = doc.add_paragraph()
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
        bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), 'E0E3FF')
        pBdr.append(bot); pPr.append(pBdr)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(12)

    def img(path, width=Inches(5.8)):
        if path and Path(path).exists():
            doc.add_picture(path, width=width)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_after = Pt(14)
        else:
            p("[Screenshot not available]", italic=True, color=(180,180,180))

    # ── Cover ──────────────────────────────────────────────────────────────
    for _ in range(5): doc.add_paragraph()

    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r  = pr.add_run("Accessories Tracking System")
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = RGBColor(30,30,46)

    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r  = pr.add_run(role["name"] + " — User Guide")
    r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = RGBColor(*C)

    doc.add_paragraph()

    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r  = pr.add_run(role["tagline"])
    r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = RGBColor(100,116,139)

    doc.add_paragraph()

    # Credentials box
    tbl = doc.add_table(rows=2, cols=2); tbl.style = 'Table Grid'
    tbl.alignment = 1  # center
    cred_data = [
        (tbl.rows[0].cells[0], "Email",    role["email"]),
        (tbl.rows[0].cells[1], "Password", role["password"]),
        (tbl.rows[1].cells[0], "Role",     role["name"]),
        (tbl.rows[1].cells[1], "Pages",    str(len(role["pages"])) + " page(s)"),
    ]
    for cell, lbl, val in cred_data:
        lp = cell.paragraphs[0]
        rl = lp.add_run(lbl + ":  ")
        rl.font.bold = True; rl.font.size = Pt(10); rl.font.color.rgb = RGBColor(*C)
        rv = lp.add_run(val)
        rv.font.size = Pt(10); rv.font.color.rgb = RGBColor(30,30,46)

    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r  = pr.add_run(f"Prepared: {date.today().strftime('%d %B %Y')}")
    r.font.size = Pt(9); r.font.color.rgb = RGBColor(148,163,184)
    pr.paragraph_format.space_before = Pt(16)

    doc.add_page_break()

    # ── Introduction ───────────────────────────────────────────────────────
    h("About This Guide", size=16, color=C)
    p(role["tagline"], size=12, bold=True, color=(30,30,46))
    doc.add_paragraph()
    p(f"This guide shows every screen the {role['name']} can access in the "
      f"Accessories Tracking System. Each page is shown with a live order loaded "
      f"so you can see exactly what data appears.", size=11, color=(71,85,105))
    doc.add_paragraph()

    # Pages list
    h("Your Pages at a Glance", size=13, color=C, before=4)
    for i, pg in enumerate(role["pages"], 1):
        para = doc.add_paragraph(style="List Number")
        r = para.add_run(pg["title"] + " — ")
        r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(*C)
        r2 = para.add_run(pg["desc"].split(".")[0] + ".")
        r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(71,85,105)
        para.paragraph_format.space_after = Pt(5)

    doc.add_page_break()

    # ── Per-page screenshots ───────────────────────────────────────────────
    h("Page Screenshots", size=16, color=C)
    p("The following pages show each screen with a real order loaded.",
      size=11, italic=True, color=(148,163,184), after=16)

    for pg in role["pages"]:
        h(pg["title"], size=14, color=C, before=8, after=4)
        p(pg["desc"], size=11, color=(71,85,105), after=10)
        img(shots.get(pg["id"]))
        divider()

    # ── Footer page ────────────────────────────────────────────────────────
    doc.add_page_break()
    for _ in range(8): doc.add_paragraph()
    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r  = pr.add_run("Accessories Tracking System  ·  Zaber & Zubair Accessories Ltd.")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(148,163,184)
    pr = doc.add_paragraph(); pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r  = pr.add_run(f"For: {role['name']}  ·  {date.today().strftime('%d %B %Y')}  ·  Confidential")
    r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = RGBColor(200,200,210)

    doc.save(role["docx"])
    print(f"    Saved: {role['docx']}")


# ── Main ──────────────────────────────────────────────────────────────────────
async def main():
    print("\n=== ATS Role Documentation Generator ===\n")
    print("Step 1: Capturing screenshots...")
    all_shots = await capture_all()

    print("\nStep 2: Building DOCX files...")
    for role in ROLES:
        print(f"\n  Building {role['docx']}...")
        build_role_docx(role, all_shots.get(role["id"], {}))

    print("\nDone! Files created:")
    for role in ROLES:
        print(f"  {role['docx']}")

if __name__ == "__main__":
    asyncio.run(main())
