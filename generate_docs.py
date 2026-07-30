from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles helper
def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x1e, 0x1e, 0x2e)
    return p

def heading3(text):
    return doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11) if p.runs else None
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def code_line(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
    p.paragraph_format.left_indent = Inches(0.4)
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'c7d2fe')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ══════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════
title = doc.add_heading('ED Module', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
title.runs[0].font.size = Pt(28)

sub = doc.add_paragraph('Export Documentation Workspace')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

date_p = doc.add_paragraph(f'Project Documentation  ·  {datetime.date.today().strftime("%d %B %Y")}')
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.size = Pt(11)
date_p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.add_paragraph()
divider()
doc.add_paragraph()

# ══════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ══════════════════════════════════════════════
heading1('1. Project Overview')
body(
    'ED Module is a web-based Export Documentation Workspace built with PHP (XAMPP / MySQL) '
    'and vanilla JavaScript. It centralises the full lifecycle of garment export documentation — '
    'from customer profiling and order intake through to commercial invoices, packing lists, '
    'certificates of origin, and shipping challans. The application runs locally at '
    'http://localhost/ed_module/ on XAMPP installed at E:\\xampp.'
)
doc.add_paragraph()

# Stack table
heading2('Technology Stack')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Technology'
for cell in hdr:
    cell.paragraphs[0].runs[0].bold = True

rows = [
    ('Backend',    'PHP 8.x (XAMPP), PDO / MySQL'),
    ('Frontend',   'Vanilla JavaScript, CSS3 (custom design system)'),
    ('Database',   'MySQL — database name: ed_module'),
    ('Server',     'Apache (XAMPP) at E:\\xampp'),
    ('Auth',       'Session-based (PHP sessions + password_hash)'),
    ('Data sync',  'localStorage cache + MySQL as primary store via fetch() API'),
]
for layer, tech in rows:
    row = tbl.add_row().cells
    row[0].text = layer
    row[1].text = tech
doc.add_paragraph()
divider()

# ══════════════════════════════════════════════
#  2. FOLDER STRUCTURE
# ══════════════════════════════════════════════
heading1('2. Folder & File Structure')
files = [
    ('ed_module/',                      'Project root'),
    ('  includes/',                     'Shared PHP includes'),
    ('    header.php',                  'Global nav, auth bootstrap, HTML head'),
    ('    footer.php',                  'Closing HTML tags'),
    ('    auth.php',                    'Session auth helpers (currentUser, requireLogin, requireAdmin, isAdmin)'),
    ('    db.php',                      'PDO singleton — getDB()'),
    ('  pages/',                        'All user-facing pages'),
    ('    login.php',                   'Sign-in form with password_verify'),
    ('    logout.php',                  'Session destroy → redirect to login'),
    ('    dashboard.php',               'Summary dashboard'),
    ('    customer-profile.php',        'Customer Profile — list + form + print'),
    ('    item-master.php',             'Item Master page'),
    ('    users.php',                   'User Management (admin only)'),
    ('    marketing-intake.php',        'Order intake form (tab 1 of 15)'),
    ('    costing-review.php … po-status.php', '14 more document/order tabs'),
    ('  api/',                          'REST-ish JSON endpoints'),
    ('    customers.php',               'GET list / GET ?id=X / POST create'),
    ('    orders.php',                  'GET / POST / PUT orders'),
    ('    buyers.php',                  'GET / POST buyers'),
    ('    page_data.php',               'GET / POST per-page JSON blobs'),
    ('  setup/',                        'One-time database setup scripts'),
    ('    schema.sql',                  'Creates all 5 tables'),
    ('    create_users.php',            'Creates users table + seeds default admin'),
    ('    run_schema.php',              'Executes schema.sql via PDO'),
    ('  styles.css',                    'Global design system CSS'),
]
for path, desc in files:
    p = doc.add_paragraph()
    r1 = p.add_run(path.ljust(46))
    r1.font.name = 'Courier New'
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph()
divider()

# ══════════════════════════════════════════════
#  3. DATABASE SCHEMA
# ══════════════════════════════════════════════
heading1('3. Database Schema (ed_module)')
tables_info = [
    ('users',       ['id', 'name', 'email', 'password (hashed)', 'role (admin|manager|staff)', 'is_active', 'created_by', 'created_at']),
    ('customers',   ['id', 'data (JSON blob of all form fields)', 'created_at', 'updated_at']),
    ('buyers',      ['id', 'name', 'country', 'contact', 'created_at']),
    ('orders',      ['id', 'customer_id', 'order_data (JSON)', 'status', 'created_at']),
    ('order_items', ['id', 'order_id', 'item_data (JSON)', 'qty', 'price']),
    ('page_data',   ['id', 'page_key', 'order_id', 'data (JSON)', 'updated_at']),
]
for tname, cols in tables_info:
    heading3(f'Table: {tname}')
    body(', '.join(cols))

heading2('Default Admin Credentials')
bullet('Email: admin@ed.local')
bullet('Password: Admin@1234')
body('Change immediately after first login.')
doc.add_paragraph()
divider()

# ══════════════════════════════════════════════
#  4. AUTHENTICATION & USER ROLES
# ══════════════════════════════════════════════
heading1('4. Authentication & User Roles')
body(
    'The app uses PHP session-based authentication. All pages include auth.php via header.php '
    'and call requireLogin() which redirects to login.php if no session exists.'
)
doc.add_paragraph()
heading2('Roles')
roles = [
    ('admin',   'Full access. Can create/edit/deactivate users, change roles, reset passwords. '
                'Sees the ⚙ Users tab in Master Data nav.'),
    ('manager', 'Access to all Order & Documents pages and Master Data. Cannot manage users.'),
    ('staff',   'Standard access to Order & Documents pages.'),
]
for role, desc in roles:
    bullet(f' — {desc}', bold_prefix=role.upper())

heading2('Auth Helpers (includes/auth.php)')
for fn in ['currentUser() → ?array  — returns session user or null',
           'requireLogin() — redirects to login.php if not authenticated',
           'requireAdmin() — 403 if not admin',
           'isAdmin() → bool']:
    code_line(fn)
doc.add_paragraph()
divider()

# ══════════════════════════════════════════════
#  5. NAVIGATION STRUCTURE
# ══════════════════════════════════════════════
heading1('5. Navigation Structure')
body(
    'The navigation has two levels rendered in includes/header.php. '
    'The top row shows the three section buttons with flow arrows between them. '
    'A tab strip below shows the tabs for the active section.'
)
doc.add_paragraph()

heading2('Top-Level Sections (nav-section-row)')
bullet('Master Data  →  Order & Documents  →  Dashboard')
body('Each page sets $navSection = "master" | "order" | "dashboard" to highlight the correct button.')
doc.add_paragraph()

heading2('Master Data Tabs')
bullet('Customer Profile  →  Item Master  →  ⚙ Users (admin only)')

heading2('Order & Documents Tabs (15 tabs)')
order_tabs = [
    'Marketing Intake', 'Costing Review', 'PI (Proforma Invoice)', 'Marketing',
    'LC (Letter of Credit)', 'PO Status', 'Bill of Exchange', 'Commercial Invoice',
    'Packing List', 'Delivery Challan', 'Truck Challan', 'Certificate of Origin',
    "Beneficiary's Certificate", 'Forwarding', 'Challan Sheet'
]
for i, t in enumerate(order_tabs, 1):
    bullet(f'{i}. {t}')

doc.add_paragraph()
divider()

# ══════════════════════════════════════════════
#  6. CUSTOMER PROFILE PAGE
# ══════════════════════════════════════════════
heading1('6. Customer Profile Page')
body(
    'The Customer Profile is the source of truth for all buyer/customer data. '
    'No separate Buyer Master exists — all customer information lives here.'
)
doc.add_paragraph()

heading2('Features Implemented')
features = [
    'Searchable paginated list at the top (10 records per page) showing: #, Company, Type, Chairman, Mobile, Head Office, Submitted date, Action',
    'Click any row → opens detail popup modal with all fields grouped into 3 sections: Core Information, Additional Details, Selected Options (shown as badges)',
    '+ New Profile button toggles the full create form below the list',
    'Form has 2 A4-sized pages with all customer fields, checkboxes, rate tables, and signature upload areas',
    'Submit button POSTs JSON to api/customers.php via fetch()',
    'Print support: form is forced display:block in @media print even when toggled closed on screen',
    'Signature upload: FileReader API shows image preview instantly on upload',
    'Clear button resets all inputs, checkboxes, and signature images',
]
for f in features:
    bullet(f)

heading2('Form Field Mapping')
body(
    '57 checkbox labels (CB_LABELS[]) and 23 text field labels (FIELD_LABELS[]) are defined '
    'as JavaScript arrays. These map DOM input positions to human-readable field names so the '
    'database stores meaningful keys instead of "cb_9" or "field_6".'
)
doc.add_paragraph()

heading2('API Endpoint: api/customers.php')
bullet('GET /api/customers.php — returns paginated list')
bullet('GET /api/customers.php?id=X — returns single record by ID')
bullet('POST /api/customers.php — creates new customer record (JSON body)')
doc.add_paragraph()
divider()

# ══════════════════════════════════════════════
#  7. USER MANAGEMENT PAGE
# ══════════════════════════════════════════════
heading1('7. User Management Page (pages/users.php)')
body('Admin-only page accessible via the ⚙ Users tab in Master Data.')
doc.add_paragraph()

heading2('Features')
bullet('Stats bar: Total Users / Active count / Admin count')
bullet('User card grid — each card shows: colour-coded avatar (initials), name, email, role selector, status dot, join date')
bullet('Avatar colour: purple = admin, blue = manager, green = staff')
bullet('"You" tag on the current user\'s own card; cannot deactivate or change own role')
bullet('Role change: inline select → auto-submits form on change')
bullet('Deactivate / Activate toggle button on each card')
bullet('Reset Password button opens a modal')
bullet('+ New User button opens Create User modal')
bullet('Both modals close on backdrop click or ✕ button')

heading2('POST Actions')
for act in ['create — insert new user', 'toggle — flip is_active', 'change_role — update role', 'reset_password — update hashed password']:
    code_line(act)
doc.add_paragraph()
divider()

# ══════════════════════════════════════════════
#  8. DESIGN SYSTEM
# ══════════════════════════════════════════════
heading1('8. Design System (styles.css)')
body('All styles live in a single styles.css file. Key design tokens:')
doc.add_paragraph()

tokens = [
    ('Primary accent',    '#6366f1 / #4f46e5  (indigo)'),
    ('Background',        '#f8f9ff  (very light indigo tint)'),
    ('Card background',   '#ffffff  with 1.5px border #e8eaff'),
    ('Border radius',     '12–16px for cards, 8–9px for inputs, 999px for pills'),
    ('Box shadow',        '0 2px 12px rgba(99,102,241,.06) — subtle indigo tint'),
    ('Font weights',      '700–800 for headings, 600 for labels, 400 for body'),
    ('Nav active pill',   'gradient linear-gradient(135deg,#6366f1,#4f46e5) with white text'),
]
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
hdr2 = tbl2.rows[0].cells
hdr2[0].text = 'Token'
hdr2[1].text = 'Value'
for c in hdr2:
    c.paragraphs[0].runs[0].bold = True
for name, val in tokens:
    r = tbl2.add_row().cells
    r[0].text = name
    r[1].text = val

doc.add_paragraph()

heading2('Key CSS Classes')
css_classes = [
    ('.nav-section-btn',        'Top-level section pills (Master Data / Order & Documents)'),
    ('.nav-home-btn',           'Dashboard button — right-aligned gradient pill'),
    ('.nav-flow-arrow',         'Arrow → between section buttons'),
    ('.tab-flow-arrow',         'Subtle arrow → between tabs in a group'),
    ('.nav-user-bar',           'User name + role badge + Sign Out — far right of nav'),
    ('.page-tab',               'Individual tab button in a tab group'),
    ('.form-card',              'White card wrapper for page sections'),
    ('.section-head',           'Section header row with title + action button'),
    ('.primary-btn',            'Indigo gradient action button'),
    ('.ghost-btn',              'Outline secondary button'),
    ('.cp-list-table',          'Customer / data list table'),
    ('.cp-type-badge',          'Coloured badge for customer type'),
    ('.cp-modal-shell',         'Full-screen overlay for detail modals'),
    ('.cp-page',                'A4-sized form page container'),
    ('.cp-page-break',          'display:none on screen, page-break in print CSS'),
    ('.um-card',                'User management card'),
    ('.um-avatar',              'Circular initials avatar with role colour'),
    ('.um-stats / .um-stat-card', 'Stats bar at top of User Management page'),
]
for cls, desc in css_classes:
    p = doc.add_paragraph()
    r1 = p.add_run(cls.ljust(32))
    r1.font.name = 'Courier New'
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x4f, 0x46, 0xe5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

doc.add_paragraph()
divider()

# ══════════════════════════════════════════════
#  9. KNOWN ISSUES FIXED
# ══════════════════════════════════════════════
heading1('9. Issues Fixed During Development')
issues = [
    ('Modal showing "CB 9 = true"',
     'Old submit code saved extraData["cb_"+i]. Fixed by defining FIELD_LABELS[] and CB_LABELS[] arrays to map DOM index → human label.'),
    ('Blank space on Customer Profile form',
     'Two causes: (1) .cp-page-num had position:absolute;bottom:20px forcing page div very tall — changed to text-align:right;margin-top:16px. '
     '(2) min-height:900px on .cp-page — removed from screen CSS, kept only in @media print.'),
    ('Print showing blank page',
     'section.form-card:first-of-type selector in print CSS hid everything. Also #cpFormSection is display:none when toggled. '
     'Fixed: added #cpFormSection { display:block !important } and matching section rule to print CSS.'),
    ('XAMPP path errors',
     'XAMPP is installed at E:\\xampp (not C:\\xampp). PHP CLI commands use & "E:\\xampp\\php\\php.exe".'),
    ('Rate table inputs invisible',
     'Added border-bottom:1.5px solid #aaa to .cp-td-input class so inputs are visible inside table cells.'),
    ('Signature upload not showing',
     'Implemented FileReader API in loadSig() — reads file as DataURL and sets it as <img> src immediately on input change.'),
]
for title_text, detail in issues:
    heading3(title_text)
    body(detail)

doc.add_paragraph()
divider()

# ══════════════════════════════════════════════
#  10. PENDING / NEXT STEPS
# ══════════════════════════════════════════════
heading1('10. Pending & Next Steps')
pending = [
    'Item Master — create/edit form with product line → line item → grade → paper combination cascading dropdowns',
    'Order & Documents pages — wire all 15 tabs to database via api/page_data.php',
    'Dashboard — live stats from DB (orders count, customers count, recent activity)',
    'Print layout verification — confirm Customer Profile 2-page print fits correctly on A4',
    'Email / notification system for order status changes (optional)',
    'Backup / export — CSV or Excel export of customer list and order data',
]
for p_text in pending:
    bullet(p_text)

doc.add_paragraph()

# Footer note
footer_p = doc.add_paragraph(f'Generated on {datetime.date.today().strftime("%d %B %Y")}  ·  ED Module v1.0  ·  Confidential')
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.runs[0].font.size = Pt(9)
footer_p.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

# Save
out = r'e:\xampp\htdocs\ed_module\ED_Module_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
